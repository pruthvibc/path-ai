import os
import json
import re
import io
import requests
import PyPDF2
from fastapi import FastAPI, Request, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv
from pypdf import PdfReader
from typing import Annotated, Optional

# 1. SETUP & CONFIGURATION
load_dotenv()
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    print("CRITICAL ERROR: GROQ_API_KEY not found.")

groq_client = Groq(api_key=api_key)
MEMORY_FILE = "memory.json"

SERPER_API_KEY = os.environ.get("SERPER_API_KEY")

SESSION_STORE: dict[str, dict] = {}


def load_memories():
    try:
        if os.path.exists(MEMORY_FILE):
            with open(MEMORY_FILE, "r") as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
                return data.get("_legacy", [])
        return []
    except Exception:
        return []

def save_memory(new_insight):
    try:
        mems = load_memories()
        clean = new_insight.strip()
        if "CANDIDATE_NAME:" in clean:
            mems = [m for m in mems if "CANDIDATE_NAME:" not in m]
        if clean and clean not in mems:
            mems.append(clean)
        full = _read_full_store()
        full["_legacy"] = mems[-30:]
        _write_full_store(full)
    except Exception as e:
        print(f"Memory Save Error: {e}")


def _read_full_store() -> dict:
    try:
        if os.path.exists(MEMORY_FILE):
            with open(MEMORY_FILE, "r") as f:
                data = json.load(f)
            if isinstance(data, list):
                return {"_legacy": data, "users": {}}
            return data
    except Exception:
        pass
    return {"_legacy": [], "users": {}}

def _write_full_store(data: dict):
    tmp = MEMORY_FILE + ".tmp"
    try:
        with open(tmp, "w") as f:
            json.dump(data, f, indent=2)
        os.replace(tmp, MEMORY_FILE)
    except Exception as e:
        print(f"Store write error: {e}")
        try:
            os.unlink(tmp)
        except:
            pass

def _make_user_id(name: str) -> str:
    import re as _re
    return _re.sub(r'[^a-z0-9]+', '_', name.lower()).strip('_')[:40]

def _hash_password(password: str) -> str:
    import hashlib
    salted = f"ai_advisor_salt_{password}_2025"
    return hashlib.sha256(salted.encode()).hexdigest()

def _check_password(password: str, stored_hash: str) -> bool:
    return _hash_password(password) == stored_hash

def load_user(user_id: str) -> dict | None:
    store = _read_full_store()
    return store.get("users", {}).get(user_id)

def save_user(user_id: str, user_data: dict):
    store = _read_full_store()
    if "users" not in store:
        store["users"] = {}
    store["users"][user_id] = user_data
    _write_full_store(store)

def upsert_user(name: str, **kwargs) -> dict:
    from datetime import datetime, timezone
    user_id = _make_user_id(name)
    existing = load_user(user_id)
    now = datetime.now(timezone.utc).isoformat()
    if existing:
        existing["last_seen"] = now
        existing["is_new"] = False
        for k, v in kwargs.items():
            if v is not None:
                existing[k] = v
        save_user(user_id, existing)
        return existing
    else:
        user = {
            "id":              user_id,
            "name":            name,
            "is_new":          True,
            "created_at":      now,
            "last_seen":       now,
            "resume_summary":  kwargs.get("resume_summary", ""),
            "gaps":            kwargs.get("gaps", []),
            "verified_skills": kwargs.get("verified_skills", []),
            "roadmap":         kwargs.get("roadmap", []),
        }
        save_user(user_id, user)
        return user


class ChatRequest(BaseModel):
    messages: list
    topic_context: Optional[str] = None
    syllabus: Optional[list] = None
    user_id: Optional[str] = None

class MatchRequest(BaseModel):
    job_description: str
    job_title: str
    job_url: Optional[str] = None
    session_id: Optional[str] = None


async def extract_text_from_pdf(file: UploadFile):
    try:
        await file.seek(0)
        reader = PdfReader(file.file)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read PDF: {e}")


def analyze_with_llm(resume_text, jd_text):
    prompt = f"""
    You are a Technical Recruiter. Analyze the Resume and JD provided.

    TASK:
    1. Extract the candidate's actual name from the resume.
    2. Identify technical skill gaps — skills required by the JD that are NOT present in the resume.

    STRICT RULE: IGNORE celebrity names. Focus on the identity as a Software Engineer.

    RESUME: {resume_text[:4000]}
    JD: {jd_text[:4000]}

    RETURN ONLY JSON:
    {{
        "name": "Extracted Name",
        "skills_found": ["A", "B"],
        "skills_missing": ["C", "D"],
        "reasoning": "Explanation"
    }}
    """
    completion = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
        temperature=0.1,
    )
    raw_response = completion.choices[0].message.content.strip()
    clean_json = re.sub(r'^```json\s*|```$', '', raw_response, flags=re.MULTILINE).strip()
    try:
        return json.loads(clean_json)
    except:
        return {"name": "Candidate", "skills_found": [], "skills_missing": [], "reasoning": "Error"}


@app.post("/api/analyze-gap")
async def analyze_gap(
    resume: Annotated[UploadFile, File(...)],
    jd:     Annotated[UploadFile, File(...)],
    user_id: Optional[str] = Form(None),
):
    resume_text = await extract_text_from_pdf(resume)
    jd_text     = await extract_text_from_pdf(jd)
    gap_data = analyze_with_llm(resume_text, jd_text)

    if not user_id:
        save_memory(f"CANDIDATE_NAME: {gap_data.get('name', 'Candidate')}")
        for skill in gap_data.get("skills_missing", []):
            save_memory(f"Gap Identified: {skill}")

    if user_id:
        try:
            user = load_user(user_id)
            if user:
                user["gaps"]           = gap_data.get("skills_missing", [])
                user["resume_summary"] = resume_text[:500]
                user["roadmap"] = []
                save_user(user_id, user)
            else:
                upsert_user(
                    gap_data.get('name', 'Candidate'),
                    resume_summary=resume_text[:500],
                    gaps=gap_data.get("skills_missing", []),
                    roadmap=[],
                )
        except Exception as e:
            print(f"User gap save error (non-fatal): {e}")
    else:
        try:
            upsert_user(
                gap_data.get('name', 'Candidate'),
                resume_summary=resume_text[:500],
                gaps=gap_data.get("skills_missing", []),
            )
        except Exception as e:
            print(f"User upsert error (non-fatal): {e}")

    return gap_data


@app.get("/api/generate-roadmap")
async def generate_roadmap(user_id: Optional[str] = None):
    try:
        if user_id:
            user = load_user(user_id)
            if user and user.get("gaps"):
                raw_gaps = user["gaps"]
                gaps = [f"Gap Identified: {g}" for g in raw_gaps]
            else:
                return {"roadmap": [], "reasoning_trace": "No gaps found for this user."}
        else:
            from urllib.parse import unquote_plus
            mems = load_memories()
            gaps = [unquote_plus(m) for m in mems if "Gap" in m]

        if not gaps:
            return {"roadmap": [], "reasoning_trace": "No gaps found."}

        gap_list = [g.replace("Gap Identified: ", "").strip() for g in gaps]

        prompt = f"""
        Based on these specific technical skill gaps: {gap_list}, create a 3-step learning roadmap.

        IMPORTANT: The roadmap must be specifically tailored to THESE gaps only.
        Do NOT create a generic roadmap. Every step must address one or more of the listed gaps.

        STRICT RULES FOR RESOURCES:
        - Only recommend from these 4 platforms: Udemy, YouTube, Coursera, NPTEL.
        - For each resource provide a relevant course or video title specific to the skill.
        - Build the link as a SEARCH URL using the skill topic as the query:
            Udemy:    https://www.udemy.com/courses/search/?q=TOPIC
            YouTube:  https://www.youtube.com/results?search_query=TOPIC
            Coursera: https://www.coursera.org/search?query=TOPIC
            NPTEL:    https://nptel.ac.in/course.html
        - Replace TOPIC with the URL-encoded skill name.
        - Provide exactly 3 resources per step, from different platforms where possible.

        RETURN ONLY A JSON LIST with no extra text:
        [{{
            "skill": "...",
            "topic": "...",
            "syllabus": ["subtopic1", "subtopic2"],
            "resources": [
                {{"platform": "YouTube", "title": "Relevant title", "url": "https://www.youtube.com/results?search_query=..."}},
                {{"platform": "Udemy",   "title": "Relevant title", "url": "https://www.udemy.com/courses/search/?q=..."}},
                {{"platform": "Coursera","title": "Relevant title", "url": "https://www.coursera.org/search?query=..."}}
            ],
            "effort": "X weeks",
            "reasoning": "..."
        }}]
        """
        comp = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.5,
        )
        clean_json = re.sub(r'^```json\s*|```$', '', comp.choices[0].message.content.strip(), flags=re.MULTILINE).strip()
        roadmap_data = json.loads(clean_json)

        if user_id:
            try:
                user = load_user(user_id)
                if user:
                    user["roadmap"] = roadmap_data
                    save_user(user_id, user)
            except Exception as e:
                print(f"Roadmap user save error (non-fatal): {e}")
        else:
            try:
                mems2 = load_memories()
                cname = next((m.split(": ")[1] for m in mems2 if "CANDIDATE_NAME" in m), None)
                if cname:
                    upsert_user(cname, roadmap=roadmap_data)
            except Exception as e:
                print(f"Roadmap legacy save error (non-fatal): {e}")

        return {"roadmap": roadmap_data}

    except Exception as e:
        return {"error": str(e)}


@app.post("/api/chat")
async def chat_with_ai(request: ChatRequest):
    try:
        user_id = request.user_id
        candidate_name = "Candidate"

        if user_id:
            user = load_user(user_id)
            if user:
                candidate_name = user.get("name", "Candidate")
        else:
            mems = load_memories()
            candidate_name = next(
                (m.split(": ")[1] for m in mems if "CANDIDATE_NAME" in m),
                "Candidate"
            )

        user_msgs = [m for m in request.messages if isinstance(m, dict) and m.get('role') == 'user']
        q_num = len(user_msgs) + 1

        target_topic   = request.topic_context if request.topic_context else "Software Engineering"
        syllabus_items = request.syllabus if request.syllabus else []
        syllabus_context = (
            f"The module syllabus covers these specific sub-skills: {', '.join(syllabus_items)}. "
            f"Rotate questions across ALL of these sub-skills — do not repeat the same sub-skill twice."
            if syllabus_items else ""
        )

        if q_num > 10:
            eval_prompt = f"""
            The 10-question technical interview for {target_topic} is complete.
            Review the interaction for {candidate_name}.
            CRITERIA: 8/10 correct = PASS.
            RETURN ONLY JSON: {{ "text": "Feedback string", "is_complete": true, "passed": true/false }}
            HISTORY: {request.messages}
            """
            comp = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": eval_prompt}],
                model="llama-3.3-70b-versatile"
            )
            content = re.sub(r'^```json\s*|```$', '', comp.choices[0].message.content.strip(), flags=re.MULTILINE).strip()
            result = json.loads(content)

            if result.get("passed"):
                if user_id:
                    try:
                        user = load_user(user_id)
                        if user:
                            verified = user.get("verified_skills", [])
                            if target_topic not in verified:
                                verified.append(target_topic)
                            gaps = user.get("gaps", [])
                            gaps = [g for g in gaps if g.lower() != target_topic.lower()]
                            user["verified_skills"] = verified
                            user["gaps"] = gaps
                            save_user(user_id, user)
                    except Exception as e:
                        print(f"Verified skill save error (non-fatal): {e}")
                else:
                    save_memory(f"VERIFIED_MASTERY: {target_topic}")
                    try:
                        mems2 = load_memories()
                        cname = next((m.split(": ")[1] for m in mems2 if "CANDIDATE_NAME" in m), None)
                        if cname:
                            legacy_user = load_user(_make_user_id(cname))
                            if legacy_user:
                                verified = legacy_user.get("verified_skills", [])
                                if target_topic not in verified:
                                    verified.append(target_topic)
                                upsert_user(cname, verified_skills=verified)
                    except Exception as e:
                        print(f"Legacy verified skill save error (non-fatal): {e}")

            return result

        if q_num <= 3:
            level = "Foundational/Easy (Syntax, definitions, core concepts)"
        elif q_num <= 9:
            level = "Application-based/Moderate (Scenario handling, common logic)"
        else:
            level = "Deep Logic/Hard (Optimization, edge cases, complex problem-solving)"

        system_instr = (
            f"ROLE: Senior Technical Examiner conducting a SILENT CERTIFICATION TEST.\n"
            f"TOPIC: {target_topic}. CANDIDATE: {candidate_name}.\n"
            f"{syllabus_context}\n"
            f"STRICT EXAM RULES:\n"
            f"1. Address the candidate as {candidate_name}.\n"
            f"2. Ask exactly ONE question. This is Question {q_num} of 10.\n"
            f"3. DIFFICULTY: {level}.\n"
            f"4. Keep the question CONCISE — maximum 2 sentences.\n"
            f"5. Stay strictly professional. No pop-culture references.\n"
            f"6. CRITICAL — THIS IS AN EXAM: Do NOT give any feedback, hints, corrections, "
            f"praise, or commentary on the candidate's previous answer. "
            f"Immediately ask the next question without reacting to their response. "
            f"Silence on correctness is mandatory until the exam ends."
        )

        messages = [{"role": "system", "content": system_instr}]
        for msg in request.messages:
            if isinstance(msg, dict) and msg.get('role') in ('user', 'assistant') and msg.get('content'):
                messages.append({"role": msg["role"], "content": msg["content"]})

        completion = groq_client.chat.completions.create(
            messages=messages,
            model="llama-3.3-70b-versatile",
            temperature=0.4
        )

        return {
            "text":        completion.choices[0].message.content,
            "q_num":       q_num,
            "is_complete": False,
            "candidate":   candidate_name
        }

    except Exception as e:
        return {"text": f"Error: {str(e)}", "is_complete": False}


@app.post("/api/evolve-resume")
async def evolve_resume(
    resume:  Annotated[UploadFile, File(...)],
    user_id: Optional[str] = Form(None),
):
    try:
        resume_text = await extract_text_from_pdf(resume)

        mastered_skills = []
        candidate_name  = "Candidate"

        if user_id:
            user = load_user(user_id)
            if user:
                candidate_name  = user.get("name", "Candidate")
                mastered_skills = user.get("verified_skills", [])
        else:
            mems = load_memories()
            candidate_name  = next((m.split(": ")[1] for m in mems if "CANDIDATE_NAME" in m), "Candidate")
            mastered_skills = [m.replace("VERIFIED_MASTERY: ", "").strip() for m in mems if "VERIFIED_MASTERY" in m]

        if not mastered_skills:
            return {"error": "No verified skills found. Complete at least one certification exam first."}

        prompt = f"""
        You are an expert resume writer. Below is a candidate's current resume text and a list of skills they have recently been AI-certified in.

        CANDIDATE NAME: {candidate_name}
        RESUME TEXT:
        {resume_text[:5000]}

        NEWLY VERIFIED SKILLS (AI-certified, must be added):
        {json.dumps(mastered_skills)}

        TASK:
        Produce a complete, improved version of this resume with these changes:
        1. Add all verified skills to the Skills section (create one if it doesn't exist).
        2. For each verified skill, add ONE strong, quantified bullet point under the most relevant work experience entry.
        3. Add or update an "AI-Verified Certifications" section at the bottom listing each verified skill with the text "Certified via AI Proctored Exam".
        4. Keep everything else exactly the same — do not invent new jobs, degrees, or facts.
        5. Return the full resume content structured for a professional document.

        RETURN ONLY JSON in this exact format:
        {{
            "candidate_name": "...",
            "summary": "one paragraph professional summary",
            "skills": ["skill1", "skill2", "...all skills including new ones"],
            "experience": [
                {{
                    "title": "Job Title",
                    "company": "Company Name",
                    "duration": "Start – End",
                    "bullets": ["bullet 1", "bullet 2", "..."]
                }}
            ],
            "education": [
                {{
                    "degree": "...",
                    "institution": "...",
                    "year": "..."
                }}
            ],
            "certifications": [
                {{
                    "name": "Skill Name",
                    "issuer": "AI Proctored Exam",
                    "year": "2025"
                }}
            ],
            "projects": [
                {{
                    "name": "...",
                    "description": "...",
                    "tech": ["..."]
                }}
            ]
        }}
        """

        comp = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.2,
        )
        raw = comp.choices[0].message.content.strip()
        clean_json = re.sub(r'^```json\s*|```$', '', raw, flags=re.MULTILINE).strip()
        resume_data = json.loads(clean_json)
        return {"resume_data": resume_data, "verified_skills": mastered_skills}

    except Exception as e:
        return {"error": str(e)}


@app.get("/api/hindsight")
def get_hindsight(user_id: Optional[str] = None):
    if user_id:
        user = load_user(user_id)
        if not user:
            return {"memories": []}
        memories = []
        memories.append(f"CANDIDATE_NAME: {user.get('name', '')}")
        for gap in user.get("gaps", []):
            memories.append(f"Gap Identified: {gap}")
        for skill in user.get("verified_skills", []):
            memories.append(f"VERIFIED_MASTERY: {skill}")
        return {"memories": memories}
    else:
        return {"memories": load_memories()}


# ─────────────────────────────────────────────────────────────────────────────
# 8. PDF GENERATION ENDPOINT
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/generate-resume-docx")
async def generate_resume_pdf(request: Request):
    import io
    from fastapi.responses import StreamingResponse
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle

    try:
        body = await request.json()
        data = body.get("resume_data", {})

        INDIGO   = colors.HexColor('#4F46E5')
        AMBER    = colors.HexColor('#D97706')
        DARK     = colors.HexColor('#1E1B4B')
        GRAY     = colors.HexColor('#374151')
        LGRAY    = colors.HexColor('#6B7280')
        WHITE    = colors.white
        BG_CERT  = colors.HexColor('#FFFBEB')
        BG_BADGE = colors.HexColor('#D97706')
        BROWN    = colors.HexColor('#92400E')
        RULE_CLR = colors.HexColor('#E0E7FF')

        S = {
            'name':      ParagraphStyle('re_name',      fontName='Helvetica-Bold',    fontSize=26,  textColor=DARK,  alignment=TA_CENTER, spaceAfter=6),
            'body':      ParagraphStyle('re_body',      fontName='Helvetica',         fontSize=10,  textColor=GRAY,  spaceAfter=4,  leading=14),
            'italic':    ParagraphStyle('re_italic',    fontName='Helvetica-Oblique', fontSize=10,  textColor=GRAY,  spaceAfter=6,  leading=15),
            'bullet':    ParagraphStyle('re_bullet',    fontName='Helvetica',         fontSize=10,  textColor=GRAY,  spaceAfter=2,  leading=14, leftIndent=16),
            'job_title': ParagraphStyle('re_job_title', fontName='Helvetica-Bold',    fontSize=11,  textColor=DARK,  spaceAfter=1,  spaceBefore=10),
            'company':   ParagraphStyle('re_company',   fontName='Helvetica',         fontSize=9.5, textColor=LGRAY, spaceAfter=4),
            'cert_name': ParagraphStyle('re_cert_name', fontName='Helvetica-Bold',    fontSize=10,  textColor=BROWN, spaceAfter=1),
            'cert_sub':  ParagraphStyle('re_cert_sub',  fontName='Helvetica-Oblique', fontSize=8.5, textColor=AMBER, spaceAfter=0),
            'badge':     ParagraphStyle('re_badge',     fontName='Helvetica-Bold',    fontSize=8,   textColor=WHITE, alignment=TA_CENTER),
        }

        def sec(title, color=None):
            c = color or INDIGO
            return [
                Spacer(1, 8),
                Paragraph(title.upper(), ParagraphStyle(
                    f're_sec_{title[:6]}', fontName='Helvetica-Bold', fontSize=9.5,
                    textColor=c, spaceBefore=4, spaceAfter=3,
                )),
                HRFlowable(width='100%', thickness=0.75, color=RULE_CLR, spaceAfter=6),
            ]

        def bul(text):
            safe_text = str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return Paragraph(f'&#8226;  {safe_text}', S['bullet'])

        def safe(val, fallback=''):
            return str(val).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') if val else fallback

        story = []
        story.append(Paragraph(safe(data.get('candidate_name'), 'Resume'), S['name']))
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width='100%', thickness=2, color=INDIGO, spaceBefore=0, spaceAfter=12))
        if data.get('summary'):
            story += sec('Professional Summary')
            story.append(Paragraph(safe(data['summary']), S['italic']))

        if data.get('skills'):
            story += sec('Skills')
            skills_txt = '  |  '.join(safe(s) for s in data['skills'])
            story.append(Paragraph(skills_txt, S['body']))

        if data.get('experience'):
            story += sec('Work Experience')
            for exp in data['experience']:
                story.append(Paragraph(safe(exp.get('title')), S['job_title']))
                story.append(Paragraph(
                    f"{safe(exp.get('company'))}  |  {safe(exp.get('duration'))}",
                    S['company']
                ))
                for b in (exp.get('bullets') or []):
                    story.append(bul(b))

        if data.get('certifications'):
            story += sec('AI-Verified Certifications', AMBER)
            for cert in data['certifications']:
                cert_name = safe(cert.get('name'))
                issuer    = safe(cert.get('issuer'))
                year      = safe(cert.get('year'))
                tbl = Table(
                    [[
                        Paragraph('VERIFIED', S['badge']),
                        Paragraph(cert_name,  S['cert_name']),
                        Paragraph(f'{issuer}  |  {year}', S['cert_sub']),
                    ]],
                    colWidths=[0.8*inch, 4.2*inch, 2.25*inch]
                )
                tbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0, 0), (0, 0),   BG_BADGE),
                    ('BACKGROUND',    (1, 0), (2, 0),   BG_CERT),
                    ('ALIGN',         (0, 0), (0, 0),   'CENTER'),
                    ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING',    (0, 0), (-1, -1), 7),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING',  (0, 0), (-1, -1), 8),
                    ('BOX',           (0, 0), (-1, -1), 0.5, colors.HexColor('#FDE68A')),
                    ('LINEAFTER',     (0, 0), (0, 0),   0.5, colors.HexColor('#FDE68A')),
                    ('ROUNDEDCORNERS', [4]),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))

        if data.get('education'):
            story += sec('Education')
            for edu in data['education']:
                degree = safe(edu.get('degree'))
                inst   = safe(edu.get('institution'))
                yr     = safe(edu.get('year'))
                story.append(Paragraph(f'<b>{degree}</b>  |  {inst}  |  {yr}', S['body']))

        if data.get('projects'):
            story += sec('Projects')
            for proj in data['projects']:
                tech_parts = [safe(t) for t in (proj.get('tech') or [])]
                tech_str   = f'  [{", ".join(tech_parts)}]' if tech_parts else ''
                story.append(Paragraph(f'<b>{safe(proj.get("name"))}</b>{tech_str}', S['job_title']))
                if proj.get('description'):
                    story.append(Paragraph(safe(proj['description']), S['body']))

        buf = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75*inch, rightMargin=0.75*inch,
            topMargin=0.75*inch,  bottomMargin=0.75*inch,
        )
        pdf_doc.build(story)
        buf.seek(0)

        filename = safe(data.get('candidate_name'), 'Resume').replace(' ', '_')
        return StreamingResponse(
            buf,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{filename}_Evolved_Resume.pdf"'}
        )

    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ─────────────────────────────────────────────────────────────────────────────
# 8. PDF GENERATION ENDPOINT (unchanged)
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/generate-resume-docx")
async def generate_resume_pdf(request: Request):
    import io
    from fastapi.responses import StreamingResponse
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle

    try:
        body = await request.json()
        data = body.get("resume_data", {})

        INDIGO   = colors.HexColor('#4F46E5')
        AMBER    = colors.HexColor('#D97706')
        DARK     = colors.HexColor('#1E1B4B')
        GRAY     = colors.HexColor('#374151')
        LGRAY    = colors.HexColor('#6B7280')
        WHITE    = colors.white
        BG_CERT  = colors.HexColor('#FFFBEB')
        BG_BADGE = colors.HexColor('#D97706')
        BROWN    = colors.HexColor('#92400E')
        RULE_CLR = colors.HexColor('#E0E7FF')

        S = {
            'name':      ParagraphStyle('re_name',      fontName='Helvetica-Bold',    fontSize=26,  textColor=DARK,  alignment=TA_CENTER, spaceAfter=0),
            'body':      ParagraphStyle('re_body',      fontName='Helvetica',         fontSize=10,  textColor=GRAY,  spaceAfter=4,  leading=14),
            'italic':    ParagraphStyle('re_italic',    fontName='Helvetica-Oblique', fontSize=10,  textColor=GRAY,  spaceAfter=6,  leading=15),
            'bullet':    ParagraphStyle('re_bullet',    fontName='Helvetica',         fontSize=10,  textColor=GRAY,  spaceAfter=2,  leading=14, leftIndent=16),
            'job_title': ParagraphStyle('re_job_title', fontName='Helvetica-Bold',    fontSize=11,  textColor=DARK,  spaceAfter=1,  spaceBefore=10),
            'company':   ParagraphStyle('re_company',   fontName='Helvetica',         fontSize=9.5, textColor=LGRAY, spaceAfter=4),
            'cert_name': ParagraphStyle('re_cert_name', fontName='Helvetica-Bold',    fontSize=10,  textColor=BROWN, spaceAfter=1),
            'cert_sub':  ParagraphStyle('re_cert_sub',  fontName='Helvetica-Oblique', fontSize=8.5, textColor=AMBER, spaceAfter=0),
            'badge':     ParagraphStyle('re_badge',     fontName='Helvetica-Bold',    fontSize=8,   textColor=WHITE, alignment=TA_CENTER),
        }

        def sec(title, color=None):
            c = color or INDIGO
            return [
                Spacer(1, 12),
                Paragraph(title.upper(), ParagraphStyle(
                    f're_sec_{title[:6]}', fontName='Helvetica-Bold', fontSize=9.5,
                    textColor=c, spaceBefore=4, spaceAfter=40,
                )),
                HRFlowable(width='100%', thickness=0.75, color=RULE_CLR, spaceAfter=6),
            ]

        def bul(text):
            safe_text = str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return Paragraph(f'&#8226;  {safe_text}', S['bullet'])

        def safe(val, fallback=''):
            return str(val).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') if val else fallback

        story = []
        story.append(Paragraph(safe(data.get('candidate_name'), 'Resume'), S['name']))
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width='100%', thickness=2, color=INDIGO, spaceBefore=0, spaceAfter=12))

        if data.get('summary'):
            story += sec('Professional Summary')
            story.append(Paragraph(safe(data['summary']), S['italic']))

        if data.get('skills'):
            story += sec('Skills')
            skills_txt = '  |  '.join(safe(s) for s in data['skills'])
            story.append(Paragraph(skills_txt, S['body']))

        if data.get('experience'):
            story += sec('Work Experience')
            for exp in data['experience']:
                story.append(Paragraph(safe(exp.get('title')), S['job_title']))
                story.append(Paragraph(
                    f"{safe(exp.get('company'))}  |  {safe(exp.get('duration'))}",
                    S['company']
                ))
                for b in (exp.get('bullets') or []):
                    story.append(bul(b))

        if data.get('certifications'):
            story += sec('AI-Verified Certifications', AMBER)
            for cert in data['certifications']:
                cert_name = safe(cert.get('name'))
                issuer    = safe(cert.get('issuer'))
                year      = safe(cert.get('year'))
                tbl = Table(
                    [[
                        Paragraph('VERIFIED', S['badge']),
                        Paragraph(cert_name,  S['cert_name']),
                        Paragraph(f'{issuer}  |  {year}', S['cert_sub']),
                    ]],
                    colWidths=[0.8*inch, 4.2*inch, 2.25*inch]
                )
                tbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0, 0), (0, 0),   BG_BADGE),
                    ('BACKGROUND',    (1, 0), (2, 0),   BG_CERT),
                    ('ALIGN',         (0, 0), (0, 0),   'CENTER'),
                    ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING',    (0, 0), (-1, -1), 7),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING',  (0, 0), (-1, -1), 8),
                    ('BOX',           (0, 0), (-1, -1), 0.5, colors.HexColor('#FDE68A')),
                    ('LINEAFTER',     (0, 0), (0, 0),   0.5, colors.HexColor('#FDE68A')),
                    ('ROUNDEDCORNERS', [4]),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))

        if data.get('education'):
            story += sec('Education')
            for edu in data['education']:
                degree = safe(edu.get('degree'))
                inst   = safe(edu.get('institution'))
                yr     = safe(edu.get('year'))
                story.append(Paragraph(f'<b>{degree}</b>  |  {inst}  |  {yr}', S['body']))

        if data.get('projects'):
            story += sec('Projects')
            for proj in data['projects']:
                tech_parts = [safe(t) for t in (proj.get('tech') or [])]
                tech_str   = f'  [{", ".join(tech_parts)}]' if tech_parts else ''
                story.append(Paragraph(f'<b>{safe(proj.get("name"))}</b>{tech_str}', S['job_title']))
                if proj.get('description'):
                    story.append(Paragraph(safe(proj['description']), S['body']))

        buf = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75*inch, rightMargin=0.75*inch,
            topMargin=0.75*inch,  bottomMargin=0.75*inch,
        )
        pdf_doc.build(story)
        buf.seek(0)

        filename = safe(data.get('candidate_name'), 'Resume').replace(' ', '_')
        return StreamingResponse(
            buf,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{filename}_Evolved_Resume.pdf"'}
        )

    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}


# ─────────────────────────────────────────────────────────────────────────────
# 8b. GENERATE RESUME AS WORD (.docx)
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/generate-resume-word")
async def generate_resume_word(request: Request):
    import io
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def safe(val, fallback=''):
        return str(val).strip() if val else fallback

    def add_horizontal_rule(doc, color_hex='4F46E5', thickness=12):
        """Add a colored bottom-border line under the last paragraph as a rule."""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def add_section_heading(doc, title):
        """Indigo bold uppercase section heading with a thin rule below."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        add_horizontal_rule(doc, color_hex='E0E7FF', thickness=6)

    try:
        body = await request.json()
        data = body.get("resume_data", {})

        doc = Document()

        # ── Page margins (0.75 inch all sides) ───────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.75)
            section.right_margin  = Inches(0.75)

        # ── Default body font ─────────────────────────────────────────────
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # ── Name ──────────────────────────────────────────────────────────
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(safe(data.get('candidate_name'), 'Resume'))
        name_run.bold = True
        name_run.font.size = Pt(26)
        name_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

        # Indigo rule below the name
        add_horizontal_rule(doc, color_hex='4F46E5', thickness=18)

        # ── Professional Summary ──────────────────────────────────────────
        if data.get('summary'):
            add_section_heading(doc, 'Professional Summary')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(safe(data['summary']))
            run.italic = True
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # ── Skills ────────────────────────────────────────────────────────
        if data.get('skills'):
            add_section_heading(doc, 'Skills')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.add_run('  |  '.join(safe(s) for s in data['skills']))

        # ── Work Experience ───────────────────────────────────────────────
        if data.get('experience'):
            add_section_heading(doc, 'Work Experience')
            for exp in data['experience']:
                # Job title
                title_p = doc.add_paragraph()
                title_p.paragraph_format.space_before = Pt(6)
                title_p.paragraph_format.space_after  = Pt(1)
                t_run = title_p.add_run(safe(exp.get('title')))
                t_run.bold = True
                t_run.font.size = Pt(11)
                t_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
                # Company | Duration
                co_p = doc.add_paragraph()
                co_p.paragraph_format.space_after = Pt(3)
                co_run = co_p.add_run(f"{safe(exp.get('company'))}  |  {safe(exp.get('duration'))}")
                co_run.font.size = Pt(9.5)
                co_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                # Bullets
                for b in (exp.get('bullets') or []):
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_after = Pt(2)
                    bp.paragraph_format.left_indent = Inches(0.2)
                    bp.add_run(safe(b))

        # ── AI-Verified Certifications ────────────────────────────────────
        if data.get('certifications'):
            add_section_heading(doc, 'AI-Verified Certifications')
            for cert in data['certifications']:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                verified_run = p.add_run('✔ VERIFIED  ')
                verified_run.bold = True
                verified_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
                name_run2 = p.add_run(safe(cert.get('name')))
                name_run2.bold = True
                name_run2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
                sub_run = p.add_run(f"  |  {safe(cert.get('issuer'))}  |  {safe(cert.get('year'))}")
                sub_run.font.size = Pt(8.5)
                sub_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
                sub_run.italic = True

        # ── Education ─────────────────────────────────────────────────────
        if data.get('education'):
            add_section_heading(doc, 'Education')
            for edu in data['education']:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                deg_run = p.add_run(safe(edu.get('degree')))
                deg_run.bold = True
                p.add_run(f"  |  {safe(edu.get('institution'))}  |  {safe(edu.get('year'))}")

        # ── Projects ──────────────────────────────────────────────────────
        if data.get('projects'):
            add_section_heading(doc, 'Projects')
            for proj in data['projects']:
                tech_parts = [safe(t) for t in (proj.get('tech') or [])]
                tech_str   = f'  [{", ".join(tech_parts)}]' if tech_parts else ''
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(2)
                proj_run = p.add_run(safe(proj.get('name')))
                proj_run.bold = True
                proj_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
                if tech_str:
                    p.add_run(tech_str)
                if proj.get('description'):
                    desc_p = doc.add_paragraph()
                    desc_p.paragraph_format.space_after = Pt(2)
                    desc_p.add_run(safe(proj['description']))

        # ── Stream the .docx ──────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = safe(data.get('candidate_name'), 'Resume').replace(' ', '_')
        return StreamingResponse(
            buf,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}_Evolved_Resume.docx"'}
        )

    except Exception as e:
        import traceback
        return {"error": str(e), "trace": traceback.format_exc()}





# # ─────────────────────────────────────────────────────────────────────────────
# # 8b. GENERATE RESUME AS WORD (.docx)
# # ─────────────────────────────────────────────────────────────────────────────

# @app.post("/api/generate-resume-word")
# async def generate_resume_word(request: Request):
#     import io
#     from fastapi.responses import StreamingResponse
#     from docx import Document
#     from docx.shared import Pt, RGBColor, Inches
#     from docx.enum.text import WD_ALIGN_PARAGRAPH
#     from docx.oxml.ns import qn
#     from docx.oxml import OxmlElement

#     def safe(val, fallback=''):
#         return str(val).strip() if val else fallback

#     def add_horizontal_rule(doc, color_hex='4F46E5', thickness=12):
#         """Add a colored bottom-border line under the last paragraph as a rule."""
#         p = doc.add_paragraph()
#         pPr = p._p.get_or_add_pPr()
#         pBdr = OxmlElement('w:pBdr')
#         bottom = OxmlElement('w:bottom')
#         bottom.set(qn('w:val'), 'single')
#         bottom.set(qn('w:sz'), str(thickness))
#         bottom.set(qn('w:space'), '1')
#         bottom.set(qn('w:color'), color_hex)
#         pBdr.append(bottom)
#         pPr.append(pBdr)
#         p.paragraph_format.space_before = Pt(0)
#         p.paragraph_format.space_after  = Pt(4)
#         return p

#     def add_section_heading(doc, title):
#         """Indigo bold uppercase section heading with a thin rule below."""
#         p = doc.add_paragraph()
#         p.paragraph_format.space_before = Pt(10)
#         p.paragraph_format.space_after  = Pt(0)
#         run = p.add_run(title.upper())
#         run.bold = True
#         run.font.size = Pt(9.5)
#         run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
#         add_horizontal_rule(doc, color_hex='E0E7FF', thickness=6)

#     try:
#         body = await request.json()
#         data = body.get("resume_data", {})

#         doc = Document()

#         # ── Page margins (0.75 inch all sides) ───────────────────────────
#         for section in doc.sections:
#             section.top_margin    = Inches(0.75)
#             section.bottom_margin = Inches(0.75)
#             section.left_margin   = Inches(0.75)
#             section.right_margin  = Inches(0.75)

#         # ── Default body font ─────────────────────────────────────────────
#         style = doc.styles['Normal']
#         style.font.name = 'Calibri'
#         style.font.size = Pt(10)
#         style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

#         # ── Name ──────────────────────────────────────────────────────────
#         name_p = doc.add_paragraph()
#         name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         name_p.paragraph_format.space_after = Pt(2)
#         name_run = name_p.add_run(safe(data.get('candidate_name'), 'Resume'))
#         name_run.bold = True
#         name_run.font.size = Pt(26)
#         name_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

#         # Indigo rule below the name
#         add_horizontal_rule(doc, color_hex='4F46E5', thickness=18)

#         # ── Professional Summary ──────────────────────────────────────────
#         if data.get('summary'):
#             add_section_heading(doc, 'Professional Summary')
#             p = doc.add_paragraph()
#             p.paragraph_format.space_after = Pt(6)
#             run = p.add_run(safe(data['summary']))
#             run.italic = True
#             run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

#         # ── Skills ────────────────────────────────────────────────────────
#         if data.get('skills'):
#             add_section_heading(doc, 'Skills')
#             p = doc.add_paragraph()
#             p.paragraph_format.space_after = Pt(4)
#             p.add_run('  |  '.join(safe(s) for s in data['skills']))

#         # ── Work Experience ───────────────────────────────────────────────
#         if data.get('experience'):
#             add_section_heading(doc, 'Work Experience')
#             for exp in data['experience']:
#                 # Job title
#                 title_p = doc.add_paragraph()
#                 title_p.paragraph_format.space_before = Pt(6)
#                 title_p.paragraph_format.space_after  = Pt(1)
#                 t_run = title_p.add_run(safe(exp.get('title')))
#                 t_run.bold = True
#                 t_run.font.size = Pt(11)
#                 t_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
#                 # Company | Duration
#                 co_p = doc.add_paragraph()
#                 co_p.paragraph_format.space_after = Pt(3)
#                 co_run = co_p.add_run(f"{safe(exp.get('company'))}  |  {safe(exp.get('duration'))}")
#                 co_run.font.size = Pt(9.5)
#                 co_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
#                 # Bullets
#                 for b in (exp.get('bullets') or []):
#                     bp = doc.add_paragraph(style='List Bullet')
#                     bp.paragraph_format.space_after = Pt(2)
#                     bp.paragraph_format.left_indent = Inches(0.2)
#                     bp.add_run(safe(b))

#         # ── AI-Verified Certifications ────────────────────────────────────
#         if data.get('certifications'):
#             add_section_heading(doc, 'AI-Verified Certifications')
#             for cert in data['certifications']:
#                 p = doc.add_paragraph()
#                 p.paragraph_format.space_after = Pt(3)
#                 verified_run = p.add_run('✔ VERIFIED  ')
#                 verified_run.bold = True
#                 verified_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
#                 name_run2 = p.add_run(safe(cert.get('name')))
#                 name_run2.bold = True
#                 name_run2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
#                 sub_run = p.add_run(f"  |  {safe(cert.get('issuer'))}  |  {safe(cert.get('year'))}")
#                 sub_run.font.size = Pt(8.5)
#                 sub_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
#                 sub_run.italic = True

#         # ── Education ─────────────────────────────────────────────────────
#         if data.get('education'):
#             add_section_heading(doc, 'Education')
#             for edu in data['education']:
#                 p = doc.add_paragraph()
#                 p.paragraph_format.space_after = Pt(3)
#                 deg_run = p.add_run(safe(edu.get('degree')))
#                 deg_run.bold = True
#                 p.add_run(f"  |  {safe(edu.get('institution'))}  |  {safe(edu.get('year'))}")

#         # ── Projects ──────────────────────────────────────────────────────
#         if data.get('projects'):
#             add_section_heading(doc, 'Projects')
#             for proj in data['projects']:
#                 tech_parts = [safe(t) for t in (proj.get('tech') or [])]
#                 tech_str   = f'  [{", ".join(tech_parts)}]' if tech_parts else ''
#                 p = doc.add_paragraph()
#                 p.paragraph_format.space_before = Pt(6)
#                 p.paragraph_format.space_after  = Pt(2)
#                 proj_run = p.add_run(safe(proj.get('name')))
#                 proj_run.bold = True
#                 proj_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
#                 if tech_str:
#                     p.add_run(tech_str)
#                 if proj.get('description'):
#                     desc_p = doc.add_paragraph()
#                     desc_p.paragraph_format.space_after = Pt(2)
#                     desc_p.add_run(safe(proj['description']))

#         # ── Stream the .docx ──────────────────────────────────────────────
#         buf = io.BytesIO()
#         doc.save(buf)
#         buf.seek(0)

#         filename = safe(data.get('candidate_name'), 'Resume').replace(' ', '_')
#         return StreamingResponse(
#             buf,
#             media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#             headers={'Content-Disposition': f'attachment; filename="{filename}_Evolved_Resume.docx"'}
#         )

#     except Exception as e:
#         import traceback
#         return {"error": str(e), "trace": traceback.format_exc()}

# ─────────────────────────────────────────────────────────────────────────────
# 9. JOB RECOMMENDATION ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/upload-resume")
async def upload_resume(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
):
    try:
        raw_bytes = await file.read()
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
        text = "".join([page.extract_text() or "" for page in pdf_reader.pages])

        extraction_prompt = f"""
        You are a resume parser. From the resume below extract:
        1. The candidate's full name (first + last, exactly as written).
        2. Up to 8 specific technical skills/tools (e.g. "React", "PostgreSQL", "Docker").

        RETURN ONLY JSON (no markdown fences):
        {{
            "name": "Full Name",
            "tags": ["skill1", "skill2", "..."]
        }}

        RESUME (first 3000 chars):
        {text[:3000]}
        """
        completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": extraction_prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.1,
        )
        raw = completion.choices[0].message.content.strip()
        clean = re.sub(r'^```json\s*|```$', '', raw, flags=re.MULTILINE).strip()
        parsed = json.loads(clean)

        candidate_name = parsed.get("name", "Candidate")
        tags           = parsed.get("tags", [])

        sid = session_id or "default"
        SESSION_STORE[sid] = {
            "resume_text": text,
            "tags":        tags,
            "name":        candidate_name,
        }

        return {"status": "success", "name": candidate_name, "tags": tags}

    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.get("/api/search-jobs")
async def search_jobs(session_id: Optional[str] = None):
    sid = session_id or "default"
    session = SESSION_STORE.get(sid)

    if not session or not session.get("resume_text"):
        raise HTTPException(status_code=400, detail="Please upload a resume first.")

    resume_text    = session["resume_text"]
    existing_tags  = session.get("tags", [])
    candidate_name = session.get("name", "Candidate")

    query_prompt = f"""
    You are a job search assistant. Based on the resume below, generate the BEST possible
    Google search query to find relevant internship or fresher job openings in Bengaluru, India.

    Rules:
    - The query must include the most relevant 2-3 tech skills from the resume.
    - Include "internship" OR "fresher" to target entry-level roles.
    - Include "Bengaluru" or "Bangalore" as the location.
    - Keep the query under 12 words total.
    - Return ONLY the raw query string, nothing else.

    Resume snippet: {resume_text[:1500]}
    Already extracted tags: {existing_tags}
    """

    query_completion = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": query_prompt}],
        model="llama-3.3-70b-versatile",
        temperature=0.2,
    )
    search_query = query_completion.choices[0].message.content.strip().strip('"').strip("'")

    url = "https://google.serper.dev/search"
    payload = json.dumps({
        "q": f"{search_query} site:linkedin.com OR site:naukri.com OR site:internshala.com OR site:unstop.com OR site:foundit.in",
        "gl": "in",
        "hl": "en",
        "num": 10
    })
    headers = {"X-API-KEY": SERPER_API_KEY, "Content-Type": "application/json"}

    try:
        response = requests.post(url, headers=headers, data=payload)
        results  = response.json()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

    job_listings = []
    for item in results.get("organic", [])[:6]:
        title   = item.get("title", "")
        snippet = item.get("snippet", "")
        link    = item.get("link", "")

        full_description = snippet
        try:
            page_resp = requests.get(link, timeout=5, headers={"User-Agent": "Mozilla/5.0"})
            if page_resp.ok:
                from html.parser import HTMLParser

                class _TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self._parts = []
                        self._skip  = False
                    def handle_starttag(self, tag, attrs):
                        if tag in ("script", "style", "nav", "header", "footer"):
                            self._skip = True
                    def handle_endtag(self, tag):
                        if tag in ("script", "style", "nav", "header", "footer"):
                            self._skip = False
                    def handle_data(self, data):
                        if not self._skip:
                            stripped = data.strip()
                            if stripped:
                                self._parts.append(stripped)

                extractor = _TextExtractor()
                extractor.feed(page_resp.text[:40000])
                page_text = " ".join(extractor._parts)
                full_description = page_text[:1200] if len(page_text) > len(snippet) else snippet
        except Exception:
            pass

        company = ""
        if " at " in title:
            company = title.split(" at ")[-1].strip()
        elif " - " in title:
            company = title.split(" - ")[-1].strip()
        else:
            company = snippet.split(".")[0].strip()

        job_listings.append({
            "title":       title,
            "company":     company,
            "link":        link,
            "snippet":     snippet,
            "description": full_description,
        })

    return {
        "jobs":            job_listings,
        "search_query":    search_query,
        "candidate_name":  candidate_name,
    }


@app.post("/api/match-job")
async def match_job(request: MatchRequest):
    sid = request.session_id or "default"
    session = SESSION_STORE.get(sid)

    if not session or not session.get("resume_text"):
        return {"error": "No resume found. Please upload your resume first."}

    resume_text = session["resume_text"]

    match_prompt = (
        "You are a Technical Recruiter. Compare the Resume against the Job Posting below.\n"
        "Be specific — identify exact matching skills and exact missing skills.\n"
        "Return ONLY a valid JSON object with these keys:\n"
        "  match_percentage (int 0-100),\n"
        "  matched_skills   (list of skills the candidate HAS that the job needs),\n"
        "  missing_skills   (list of skills the job needs that the candidate LACKS),\n"
        "  verdict          (2-sentence honest assessment),\n"
        "  improvement_advice (1 actionable sentence the candidate can act on today).\n\n"
        f"RESUME (full text):\n{resume_text[:4000]}\n\n"
        f"JOB POSTING:\nTitle: {request.job_title}\n{request.job_description[:3000]}"
    )

    try:
        completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system",  "content": "Return a valid JSON object only. No markdown, no explanation."},
                {"role": "user",    "content": match_prompt}
            ],
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            temperature=0.1,
        )
        return json.loads(completion.choices[0].message.content)
    except Exception as e:
        return {"error": f"Analysis failed: {str(e)}"}


# ─────────────────────────────────────────────────────────────────────────────
# 10. AUTH ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

class AuthRequest(BaseModel):
    username: str
    password: str

@app.post("/api/signup")
def signup(req: AuthRequest):
    username = req.username.strip()
    password = req.password.strip()

    if not username or len(username) < 2:
        return {"success": False, "error": "Username must be at least 2 characters."}
    if not password or len(password) < 4:
        return {"success": False, "error": "Password must be at least 4 characters."}

    user_id  = _make_user_id(username)
    existing = load_user(user_id)

    if existing:
        return {"success": False, "error": f"Username '{username}' is already taken. Please choose a different one."}

    from datetime import datetime, timezone
    now  = datetime.now(timezone.utc).isoformat()
    user = {
        "id":              user_id,
        "name":            username,
        "password_hash":   _hash_password(password),
        "is_new":          True,
        "created_at":      now,
        "last_seen":       now,
        "resume_summary":  "",
        "gaps":            [],
        "verified_skills": [],
        "roadmap":         [],
    }
    save_user(user_id, user)

    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {"success": True, "is_new": True, "user": safe_user}


@app.post("/api/login")
def login(req: AuthRequest):
    username = req.username.strip()
    password = req.password.strip()

    if not username or not password:
        return {"success": False, "error": "Please enter both username and password."}

    user_id = _make_user_id(username)
    user    = load_user(user_id)

    if not user:
        return {"success": False, "error": f"No account found for '{username}'. Please sign up first."}

    if not _check_password(password, user.get("password_hash", "")):
        return {"success": False, "error": "Incorrect password. Please try again."}

    from datetime import datetime, timezone
    user["last_seen"] = datetime.now(timezone.utc).isoformat()
    user["is_new"]    = False
    save_user(user_id, user)

    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {"success": True, "is_new": False, "user": safe_user}


@app.get("/api/user/{user_id}")
def get_user(user_id: str):
    user = load_user(user_id)
    if not user:
        return {"exists": False, "user": None}
    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {"exists": True, "user": safe_user}


@app.get("/api/check-user")
def check_user(name: str):
    if not name or not name.strip():
        return {"exists": False, "user": None}
    user_id = _make_user_id(name.strip())
    user    = load_user(user_id)
    if not user:
        return {"exists": False, "user": None}
    from datetime import datetime, timezone
    user["last_seen"] = datetime.now(timezone.utc).isoformat()
    user["is_new"]    = False
    save_user(user_id, user)
    safe_user = {k: v for k, v in user.items() if k != "password_hash"}
    return {"exists": True, "user": safe_user}


@app.get("/api/all-users")
def get_all_users():
    store = _read_full_store()
    users = [
        {k: v for k, v in u.items() if k != "password_hash"}
        for u in store.get("users", {}).values()
    ]
    return {"users": users}


@app.get("/")
def health_check():
    return {"status": "Backend Active"}


# ─────────────────────────────────────────────────────────────────────────────
# SELF-INTRODUCTION COACH
# ─────────────────────────────────────────────────────────────────────────────

from pydantic import BaseModel
from typing import Optional
 
class SelfIntroRequest(BaseModel):
    intro_text: str
    user_id: Optional[str] = None
    target_role: Optional[str] = None   # pulled from job recommendations if available
 
 
@app.post("/api/analyze-intro")
async def analyze_self_intro(request: SelfIntroRequest):
    """
    Analyzes a candidate's self-introduction and returns:
    - Effectiveness score (0-100)
    - Sub-scores: clarity, relevance, impact, professionalism
    - Mistake list with labels and categories
    - AI-improved version of the intro
    - Personalized tips based on resume data (if user_id provided)
    - What-to-include checklist status
    """
    try:
        intro_text = request.intro_text.strip()
        if not intro_text or len(intro_text) < 20:
            return {"error": "Please provide a meaningful self-introduction (at least a few sentences)."}
 
        # ── Pull user resume context for personalization ──────────────────────
        resume_context = ""
        candidate_name = "Candidate"
        target_role = request.target_role or "Software Engineering / Tech"
 
        if request.user_id:
            user = load_user(request.user_id)
            if user:
                candidate_name = user.get("name", "Candidate")
                resume_summary = user.get("resume_summary", "")
                verified_skills = user.get("verified_skills", [])
                gaps = user.get("gaps", [])
                if resume_summary:
                    resume_context = f"""
CANDIDATE PROFILE (from uploaded resume):
- Name: {candidate_name}
- Resume Summary: {resume_summary[:600]}
- Verified Skills: {', '.join(verified_skills) if verified_skills else 'None yet'}
- Known Gaps: {', '.join(gaps[:5]) if gaps else 'None identified'}
"""
 
        prompt = f"""
You are an elite placement coach at a top engineering college in India. Analyze this student's self-introduction for tech internship/job interviews.
 
{resume_context}
TARGET ROLE: {target_role}
 
SELF-INTRODUCTION TO ANALYZE:
\"\"\"{intro_text}\"\"\"
 
Return ONLY a valid JSON object with this EXACT structure (no markdown, no explanation):
 
{{
  "effectiveness_score": <integer 0-100 — compute using SCORING RULES below>,
  "sub_scores": {{
    "clarity":         <integer 1-10>,
    "relevance":       <integer 1-10>,
    "impact":          <integer 1-10>,
    "professionalism": <integer 1-10 — give 9-10 if greeting AND closing present, 7-8 if one is present, 5-6 if neither>
  }},
  "estimated_duration_seconds": <integer, estimated spoken time at normal pace>,
  "mistakes": [
    {{
      "text":     "<shortest exact phrase from intro that is the problem, or 'Missing' if absent>",
      "issue":    "<one of: 'Missing Greeting' | 'Missing Thank You' | 'Generic Claim' | 'No Evidence' | 'Weak Closing' | 'Filler Words' | 'Grammar Error' | 'Irrelevant Info' | 'Robotic Tone' | 'Exaggeration' | 'Lack of Specificity'>",
      "severity": "<'high' | 'medium' | 'low'>",
      "fix":      "<one actionable fix>"
    }}
  ],
  "key_issue": "<2 sentences: the single most impactful thing to improve>",
  "what_included": {{
    "proper_greeting":          <true if starts with Good morning/Hi/Hello/Good afternoon — false otherwise>,
    "name_and_academic":        <true if name + college + year/semester/CGPA present — false otherwise>,
    "core_technical_skills":    <true if 3+ specific tech skills named — false otherwise>,
    "impactful_projects":       <true if at least one project described — false otherwise>,
    "experience_highlights":    <true if any internship/hackathon/work experience OR quantified project metrics mentioned — false otherwise>,
    "key_strengths_with_proof": <true if a strength is backed by a specific example or number — false otherwise>,
    "career_goal":              <true if a clear career direction or target role stated — false otherwise>,
    "proper_closing":           <true if ends with Thank you / I look forward to / That's all about me — false otherwise>
  }},
  "improved_intro": "<Rewritten version. MUST: (1) open with 'Good morning sir/ma'am, thank you for this opportunity', (2) keep the candidate's real name, college, skills and projects from the original, (3) add 1-2 specific metrics if missing, (4) end with 'Thank you for your time. I look forward to contributing to your team.' Target 45-55 seconds spoken. Sound natural, not robotic.>",
  "improvement_highlights": [
    "<change1 starting with action verb>",
    "<change2>",
    "<change3>"
  ],
  "improved_skills_shown": ["<skill1>","<skill2>","<skill3>","<skill4>","<skill5>","<skill6>"],
  "personalized_tips": [
    "<specific actionable tip for this candidate>",
    "<another tip>",
    "<another tip>"
  ],
  "readiness_verdict": "<'Not Ready' | 'Needs Work' | 'Almost There' | 'Interview Ready'>",
  "readiness_color":   "<'red' | 'orange' | 'yellow' | 'green'>"
}}
 
═══════════════════════════════════════════════
SCORING RULES — compute effectiveness_score exactly like this:
═══════════════════════════════════════════════
Step 1 — CHECKLIST POINTS (what_included):
  proper_greeting          = true  → +15 pts  (very important in Indian interviews)
  name_and_academic        = true  → +10 pts
  core_technical_skills    = true  → +10 pts
  impactful_projects       = true  → +10 pts
  experience_highlights    = true  → +8  pts
  key_strengths_with_proof = true  → +8  pts
  career_goal              = true  → +8  pts
  proper_closing           = true  → +15 pts  (very important in Indian interviews)
  ──────────────────────────────────────────
  Max checklist total = 84 pts
 
Step 2 — MISTAKE DEDUCTIONS:
  Each 'high'   severity mistake: -12 pts
  Each 'medium' severity mistake: -4  pts
  Each 'low'    severity mistake: -1  pts
 
Step 3 — BONUS:
  If average of 4 sub_scores >= 8.0: +8 pts
  If average of 4 sub_scores >= 6.5: +4 pts
 
Step 4 — Clamp final score to range [15, 100].
 
Step 5 — Map to verdict:
  0–45  = Not Ready      (red)
  46–65 = Needs Work     (orange)
  66–79 = Almost There   (yellow)
  80–100= Interview Ready(green)
 
═══════════════════════════════════════════════
MISTAKE SEVERITY GUIDE:
═══════════════════════════════════════════════
HIGH severity (each costs -12):
  - No greeting at all (Missing Greeting)
  - No closing/thank-you at all (Missing Thank You)
  - Starts with "Myself" (Grammar Error)
  - Mentions parents/family (Irrelevant Info)
 
MEDIUM severity (each costs -4):
  - Claims a strength with zero proof (Generic Claim)
  - Project mentioned but zero metrics or impact (No Evidence)
  - Career goal is vague ("I want to grow") (Weak Closing)
  - Filler sentences that add no value (Filler Words)
 
LOW severity (each costs -1):
  - Minor grammar issues
  - Slightly robotic phrasing
  - Minor lack of specificity
 
IMPORTANT — DO NOT flag these as mistakes:
  - Mentioning college name, CGPA, or year of study → this is REQUIRED in Indian interviews, never flag as Irrelevant Info
  - Having a career goal statement → never flag as Poor Ending
  - Describing projects in detail → never flag as Too Long unless > 90 seconds
 
CALIBRATION EXAMPLES:
  - Intro with greeting + all 8 checklist items + only medium mistakes → score 75-85
  - Intro with greeting + closing + good projects + metrics → score 82-90
  - Intro missing greeting AND closing, generic claims → score 45-58
  - Perfect intro: greeting + 3 projects with metrics + closing → score 88-95
  - AI-rewritten improved intro should score 85-92
 
mistakes array: find 2-5 real issues only. Do not fabricate problems to fill the array.
improved_intro: keep the candidate's real name, college, and actual projects — only improve phrasing and add missing elements.
"""
 
        completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a strict placement coach. Return only valid JSON. No markdown fences."},
                {"role": "user", "content": prompt}
            ],
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            temperature=0.3,
        )
 
        result = json.loads(completion.choices[0].message.content)
 
        # ── Ensure sub_scores always exists with defaults ──────────────────
        if "sub_scores" not in result or not isinstance(result.get("sub_scores"), dict):
            result["sub_scores"] = {"clarity": 5, "relevance": 5, "impact": 5, "professionalism": 5}
        else:
            for key in ["clarity", "relevance", "impact", "professionalism"]:
                result["sub_scores"].setdefault(key, 5)
 
        # ── Ensure new what_included keys always exist ─────────────────────
        if "what_included" in result and isinstance(result["what_included"], dict):
            result["what_included"].setdefault("proper_greeting", False)
            result["what_included"].setdefault("proper_closing", False)
 
        # ── Persist intro score to user record ────────────────────────────────
        if request.user_id:
            try:
                user = load_user(request.user_id)
                if user:
                    user["last_intro_score"] = result.get("effectiveness_score", 0)
                    user["last_intro_verdict"] = result.get("readiness_verdict", "")
                    save_user(request.user_id, user)
            except Exception as e:
                print(f"Intro score save error (non-fatal): {e}")
 
        return result
 
    except json.JSONDecodeError as e:
        return {"error": f"AI response parse error: {str(e)}"}
    except Exception as e:
        return {"error": f"Analysis failed: {str(e)}"}
 
 
@app.post("/api/regenerate-intro")
async def regenerate_intro(request: SelfIntroRequest):
    """
    Given feedback context, regenerates an improved intro with different phrasing.
    Useful for the 'Try Again' flow on the frontend.
    """
    try:
        resume_context = ""
        candidate_name = "Candidate"
        target_role = request.target_role or "Software Engineering"
 
        if request.user_id:
            user = load_user(request.user_id)
            if user:
                candidate_name = user.get("name", "Candidate")
                resume_context = f"Resume snippet: {user.get('resume_summary', '')[:400]}"
                target_role = target_role or "Software Engineering"
 
        prompt = f"""
You are a placement coach. Rewrite this self-introduction to be placement-ready for a {target_role} role.
{resume_context}
 
ORIGINAL:
\"\"\"{request.intro_text}\"\"\"
 
Rules:
- 45-60 seconds when spoken (~120-150 words)
- MUST start with a polite greeting: "Good morning sir/ma'am, thank you for this opportunity" or "Hello, I'm [name]"
- Include: name, college, semester, 1-2 projects with metrics, 2-3 tech skills, career goal
- Natural, confident tone — not robotic
- MUST end with a polite closing: "Thank you for your time" or "I look forward to being part of your team"
 
Return ONLY JSON: {{"improved_intro": "...", "word_count": <int>}}
"""
        comp = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"},
            temperature=0.7,
        )
        return json.loads(comp.choices[0].message.content)
    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)